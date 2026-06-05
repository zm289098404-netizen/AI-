"""标书导出：Markdown -> Word(docx) / PDF（中文友好）。"""
import io
import re


def _inline_bold_segments(text: str):
    """将 **粗体** 拆分为 (segment, is_bold) 列表。"""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    out = []
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            out.append((p[2:-2], True))
        else:
            out.append((p, False))
    return out


def _iter_lines(markdown: str):
    for raw in markdown.replace("\r\n", "\n").split("\n"):
        yield raw.rstrip()


# ---------------- Word ----------------
def to_docx(title: str, markdown: str) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    # 设置默认中文字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    try:
        style._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
            "微软雅黑",
        )
    except Exception:
        pass

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in _iter_lines(markdown):
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            for seg, bold in _inline_bold_segments(line[2:]):
                r = p.add_run(seg)
                r.bold = bold
        else:
            p = doc.add_paragraph()
            for seg, bold in _inline_bold_segments(line):
                r = p.add_run(seg)
                r.bold = bold

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------- PDF ----------------
_CJK_REGISTERED = False


def _ensure_cjk_font():
    global _CJK_REGISTERED
    if _CJK_REGISTERED:
        return "STSong-Light"
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    _CJK_REGISTERED = True
    return "STSong-Light"


def _esc(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # 将 **粗体** 转为 <b>
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return text


def to_pdf(title: str, markdown: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
    )

    font = _ensure_cjk_font()
    styles = getSampleStyleSheet()

    def mk(name, size, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], fontName=font, fontSize=size, leading=size * 1.5, **kw)

    s_title = mk("zh-title", 20, alignment=1, spaceAfter=14)
    s_h1 = mk("zh-h1", 15, spaceBefore=10, spaceAfter=6, textColor=colors.HexColor("#1e3a8a"))
    s_h2 = mk("zh-h2", 13, spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#1e3a8a"))
    s_h3 = mk("zh-h3", 12, spaceBefore=6, spaceAfter=3)
    s_body = mk("zh-body", 10.5)
    s_quote = mk("zh-quote", 9.5, textColor=colors.grey, backColor=colors.HexColor("#fef3c7"),
                 borderPadding=4, leftIndent=6)
    s_bullet = mk("zh-bullet", 10.5, leftIndent=14, bulletIndent=4)

    buf = io.BytesIO()
    docp = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm, title=title,
    )
    flow = [Paragraph(_esc(title), s_title), Spacer(1, 4)]

    for line in _iter_lines(markdown):
        if not line.strip():
            flow.append(Spacer(1, 4))
            continue
        if line.startswith("### "):
            flow.append(Paragraph(_esc(line[4:]), s_h3))
        elif line.startswith("## "):
            flow.append(Paragraph(_esc(line[3:]), s_h2))
        elif line.startswith("# "):
            flow.append(Paragraph(_esc(line[2:]), s_h1))
        elif line.startswith("> "):
            flow.append(Paragraph(_esc(line[2:]), s_quote))
        elif line.startswith("- ") or line.startswith("* "):
            flow.append(Paragraph("• " + _esc(line[2:]), s_bullet))
        else:
            flow.append(Paragraph(_esc(line), s_body))

    docp.build(flow)
    return buf.getvalue()
