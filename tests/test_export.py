"""导出模块（Word / PDF）单元测试。"""
import io

from app.export import to_docx, to_pdf, _inline_bold_segments

MARKDOWN = """# 某客户投标方案

> 这是引用提示。

## 1. 项目背景
这是**加粗**正文内容。

- 要点一
- 要点二
"""


def test_inline_bold_segments():
    segs = _inline_bold_segments("普通**加粗**结尾")
    assert ("普通", False) in segs
    assert ("加粗", True) in segs
    assert ("结尾", False) in segs


def test_to_docx_produces_valid_zip():
    data = to_docx("某客户投标方案", MARKDOWN)
    assert data[:2] == b"PK"  # docx 是 zip
    assert len(data) > 1000


def test_to_docx_contains_chinese_text():
    import docx

    data = to_docx("某客户投标方案", MARKDOWN)
    doc = docx.Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert "某客户投标方案" in texts
    assert any("项目背景" in t for t in texts)
    # 标题样式被正确应用
    styles = [p.style.name for p in doc.paragraphs]
    assert any("Heading" in s for s in styles)


def test_to_pdf_produces_valid_pdf():
    data = to_pdf("某客户投标方案", MARKDOWN)
    assert data[:5] == b"%PDF-"
    assert len(data) > 1000


def test_to_pdf_extractable_chinese():
    from pypdf import PdfReader

    data = to_pdf("某客户投标方案", MARKDOWN)
    reader = PdfReader(io.BytesIO(data))
    text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
    assert "投标" in text
